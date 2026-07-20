"""
resume_parser.py
Extracts raw text from uploaded resume files (PDF / DOCX / TXT).
"""

import io
from PyPDF2 import PdfReader
import docx


def parse_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    text_chunks = []
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text_chunks.append(extracted)
    return "\n".join(text_chunks)


def parse_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also capture text inside tables (common in resumes with skill grids)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)


def parse_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def parse_resume(uploaded_file) -> str:
    """
    Accepts a Streamlit UploadedFile object and routes to the right parser
    based on file extension.
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if filename.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return parse_docx(file_bytes)
    elif filename.endswith(".txt"):
        return parse_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload PDF, DOCX, or TXT.")


def get_file_type(uploaded_file) -> str:
    filename = uploaded_file.name.lower()
    if filename.endswith(".pdf"):
        return "pdf"
    elif filename.endswith(".docx"):
        return "docx"
    elif filename.endswith(".txt"):
        return "txt"
    return "unknown"
