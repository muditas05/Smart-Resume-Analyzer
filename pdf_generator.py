"""
pdf_generator.py
Exports tailored resumes / cover letters as polished PDFs (ReportLab)
and DOCX files (python-docx).
"""

import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from docx import Document
from docx.shared import Pt


def _build_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ResumeHeading",
        fontSize=16,
        leading=20,
        spaceAfter=6,
        textColor=colors.HexColor("#1a1a2e"),
    ))
    styles.add(ParagraphStyle(
        name="ResumeBody",
        fontSize=10.5,
        leading=15,
        spaceAfter=8,
    ))
    return styles


def generate_pdf(title: str, body_text: str) -> bytes:
    """
    Renders plain text (paragraphs separated by blank lines) into a clean PDF.
    Returns raw PDF bytes suitable for a Streamlit download_button.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        leftMargin=0.85 * inch, rightMargin=0.85 * inch,
    )
    styles = _build_styles()
    story = [
        Paragraph(title, styles["ResumeHeading"]),
        HRFlowable(width="100%", color=colors.HexColor("#cccccc")),
        Spacer(1, 12),
    ]

    for para in body_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        safe_para = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        safe_para = safe_para.replace("\n", "<br/>")
        story.append(Paragraph(safe_para, styles["ResumeBody"]))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def generate_docx(title: str, body_text: str) -> bytes:
    """
    Renders plain text into a DOCX file. Returns raw bytes.
    """
    document = Document()

    heading = document.add_heading(title, level=1)

    for para in body_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = document.add_paragraph(para)
        p.style.font.size = Pt(11)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()
